"""
md_to_meeting_docx.py  ──  NLM 結構化 Markdown → 集團格式會議紀錄（三版同出）
==========================================================================
用法（命令列）：
    python md_to_meeting_docx.py input.md
    python md_to_meeting_docx.py --clip          # 從剪貼簿讀取
    python md_to_meeting_docx.py --text "..."    # 直接傳入字串

從 upload_gui 呼叫：
    from md_to_meeting_docx import convert
    paths = convert(md_text)          # → [docx, html_single, html_rwd]
"""

import os, re, sys, copy, shutil, datetime, argparse
from pathlib import Path

# ── 路徑設定 ─────────────────────────────────────────────────────────────────
_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
SAMPLE_DIR  = os.path.join(_SCRIPT_DIR, "meeting_sample")
OUTPUT_DIR  = os.path.join(_SCRIPT_DIR, "data", "meeting_export")
KEEPER_NAME = "NotebookLM 整理版 (Jennifer)"

# ── 資料類別 ──────────────────────────────────────────────────────────────────
class MeetingData:
    def __init__(self):
        self.subject    = ""
        self.date       = ""
        self.time       = ""
        self.location   = ""
        self.attendees  = ""
        self.keeper     = KEEPER_NAME
        self.agenda     = "待辦事項進度討論（報告人：各DRI）"
        self.action_items: list[tuple] = []   # (num, text, dri, due)
        self.discussions:  list[tuple] = []   # (title, [bullets])
        self.chair_instrs: list[tuple] = []   # (person, [bold_bullets])

# ── MD 解析 ───────────────────────────────────────────────────────────────────
_STRIP_MD   = re.compile(r'\*{1,3}([^*]+)\*{1,3}')
_STRIP_CITE = re.compile(r'\s*\[\d+(?:[,，]\s*\d+)*\]')   # 移除 NLM 引用標記 [1] [2,3]

def _clean(s: str) -> str:
    s = _STRIP_CITE.sub('', s)           # 先去引用標記
    return _STRIP_MD.sub(r'\1', s).strip()

# ── Section 分割（任何 ## 標題，key 正規化為大寫+底線）───────────────────────
def _split_sections(text: str) -> dict:
    """
    將所有 ## 標題切成 {正規化KEY: 內容} dict。
    正規化：去除前後空白、轉大寫、空格→底線、去除非字母數字符號。
    例：'## Meeting Info' → 'MEETING_INFO'
        '## 會議資訊'     → '會議資訊'（原樣保留中文）
    """
    result: dict[str, str] = {}
    current = None
    buf: list[str] = []
    for line in text.splitlines():
        m = re.match(r'^##\s+(.+)', line)
        if m:
            if current is not None:
                result[current] = '\n'.join(buf)
            raw = m.group(1).strip()
            # 正規化：英文轉大寫+底線；中文原樣
            key = re.sub(r'\s+', '_', raw).upper()
            key = re.sub(r'[^\w一-鿿]', '', key)  # 只保留字母數字底線中文
            current = key
            buf = []
        elif current is not None:
            buf.append(line)
    if current is not None:
        result[current] = '\n'.join(buf)
    return result

def _find_section(sections: dict, *candidates) -> str:
    """在 sections dict 中依候選 key 清單尋找第一個命中的內容。"""
    for c in candidates:
        if c in sections:
            return sections[c]
    # 模糊比對：找包含候選關鍵字的 key
    for c in candidates:
        for k, v in sections.items():
            if c in k:
                return v
    return ''

def parse_md(text: str) -> MeetingData:
    """
    自動判斷格式並解析，回傳 MeetingData。
    優先嘗試結構化解析（有 ## 標題），fallback 到自由格式。
    """
    md = MeetingData()
    sections = _split_sections(text)
    # 只要有任何 ## 標題就嘗試結構化解析
    if sections:
        _parse_structured(text, md, sections)
    # 若結構化結果不完整，再補充 freeform
    if not md.subject or (not md.discussions and not md.action_items):
        _parse_freeform(text, md)
    # 強制重新編號 Action Items（NLM 有時跳號，如 1,2,3,5,6）
    md.action_items = [
        (str(i + 1), text, dri, due)
        for i, (_, text, dri, due) in enumerate(md.action_items)
    ]
    return md

# ── 結構化格式解析 ────────────────────────────────────────────────────────────
def _parse_structured(text: str, md: MeetingData, sections: dict | None = None):
    if sections is None:
        sections = _split_sections(text)

    # ── MEETING_INFO ──────────────────────────────────────────────────────────
    info = _find_section(sections,
        'MEETING_INFO', 'MEETINGINFO', 'MEETING_INFORMATION',
        '會議資訊', '基本資訊', '會議基本資訊', '會議信息',
        '會議基本資料', '基本資料', '會議詳情', 'MEETING_DETAILS')

    def field(*keys) -> str:
        """從 info 區塊搜尋欄位值，支援：
        - key: value 格式（捕捉到 | 或行末為止）
        - - **key**: value 格式
        - Markdown 表格 | key | value | 格式
        info 為空時 fallback 搜尋頁首段落（第一個 ## 之前的文字）。
        """
        # 搜尋範圍優先順序：
        #   1. info 段（最精準）
        #   2. 頁首（第一個 ## 之前，含 MEETING_INFO 表格但不含討論內文）
        if info:
            search_scopes = [info]
        else:
            # 頁首 = 第一個 ## 標題之前的文字
            m_hdr = re.search(r'^##\s', text, re.MULTILINE)
            preamble = text[:m_hdr.start()] if m_hdr else text
            search_scopes = [preamble] if preamble.strip() else [text]

        def _try(scope: str) -> str:
            for key in keys:
                # 格式1: key: value（不跨 | 和換行）
                pat = (fr'(?:^|[-*])\s*\**\s*{re.escape(key)}\s*\**'
                       fr'\s*[：:]\s*([^|\n]+)')
                m = re.search(pat, scope, re.IGNORECASE | re.MULTILINE)
                if m:
                    v = _clean(m.group(1))
                    if v and '|' not in v:   # 雙重保險：排除含 | 的值
                        return v
            for key in keys:
                # 格式2: Markdown 表格  | key | value |
                pat = fr'^\s*\|\s*\**\s*{re.escape(key)}\s*\**\s*\|\s*([^|\n]+?)\s*\|'
                m = re.search(pat, scope, re.IGNORECASE | re.MULTILINE)
                if m:
                    val = _clean(m.group(1))
                    # 排除分隔行（--- 或 :---: 等）
                    if val and not re.match(r'^[-:\s＿]+$', val):
                        return val
            return ''

        for scope in search_scopes:
            v = _try(scope)
            if v:
                return v
        return ''

    md.subject   = field('subject', '主題', '會議名稱', '名稱', '標題',
                         '會議 Subject', 'Meeting Subject', 'Subject')
    md.date      = field('date', '日期', '會議日期', '時間日期', 'Date')
    md.time      = field('time', '時間', '會議時間', '起迄時間', 'Time')
    md.location  = field('location', '地點', '會議地點', '場地', '地址', 'Location')
    md.attendees = field('attendees', '出席者', '出席', '與會者', '參與者', '出席人員',
                         'Attendees', '出席人')
    keeper       = field('keeper', '記錄', '記錄人', '會議記錄人', '紀錄人',
                         'Keeper', '記錄 Keeper', 'Meeting Keeper')
    if keeper:
        md.keeper = keeper
    ag = field('agenda', '議程', '會議議程', '主要議程')
    if ag:
        md.agenda = ag

    # ── ACTION_ITEMS ──────────────────────────────────────────────────────────
    action_block = _find_section(sections,
        'ACTION_ITEMS', 'ACTIONITEMS', 'ACTION',
        '待辦事項', '行動項目', 'ACTION_ITEM')
    for line in action_block.splitlines():
        parts = [p.strip() for p in line.split('|')]
        parts = [p for p in parts if p and not re.match(r'^[-:＿]+$', p)]
        if len(parts) >= 4 and re.match(r'^\d+', parts[0]):
            md.action_items.append((parts[0], _clean(parts[1]),
                                    _clean(parts[2]), _clean(parts[3])))

    # ── DISCUSSION ────────────────────────────────────────────────────────────
    disc = _find_section(sections,
        'DISCUSSION', '討論', '討論事項', '討論事項與決議',
        'DISCUSSIONS', '議題', '會議討論')
    _parse_discussion_block(disc, md)

    # ── CHAIR_INSTRUCTIONS ────────────────────────────────────────────────────
    chair = _find_section(sections,
        'CHAIR_INSTRUCTIONS', 'CHAIRINSTRUCTIONS', 'CHAIR',
        '主管指示', '長官指示', '指示', '主席指示', '指示事項')
    _parse_chair_block(chair, md)

def _parse_discussion_block(text: str, md: MeetingData):
    cur_title, cur_bullets = None, []
    for line in text.splitlines():
        stripped = line.strip()
        # 跳過 Markdown 表格列（分隔行 |---|---| 或資料行 | x | y |）
        if stripped.startswith('|'):
            continue
        if re.match(r'^\s*#{2,3}\s+', line):   # ### 或 ## 子標題
            if cur_title is not None:
                md.discussions.append((cur_title, cur_bullets[:]))
            cur_title = _clean(re.sub(r'^\s*#{2,3}\s+', '', line))
            cur_bullets = []
        elif re.match(r'^\s*\d+\.\s+\S', line) and cur_title is None:
            # 無 ### 標題時，以 "1. 議題名稱" 作為段落標題
            cur_title = _clean(re.sub(r'^\s*\d+\.\s+', '', line))
            cur_bullets = []
        elif cur_title is not None and re.match(r'^\s*[-*•]\s+', line):
            b = _clean(re.sub(r'^\s*[-*•]\s+', '', line))
            if b:
                cur_bullets.append(b)
        elif cur_title is not None and stripped and not re.match(r'^\s*#', line):
            # 非空白、非標題行也收為要點
            b = _clean(line)
            if b and len(b) > 3:
                # 延續行偵測：前一個 bullet 若不以句末符號結尾，視為同句延續，合併
                # 注意：「1.」「2.」等數字標記末尾的「.」不算句末，用 (?<!\d) 排除
                _sent_end = re.search(
                    r'[。！？…!\?]\s*(?:\[\d+\])?\s*$'   # 中日英問嘆號
                    r'|(?<!\d)\.\s*(?:\[\d+\])?\s*$',     # 句末句點（排除數字後的點）
                    cur_bullets[-1]) if cur_bullets else None
                if (cur_bullets
                        and not _sent_end
                        and not re.match(r'^\s*[-*•]', line)):
                    cur_bullets[-1] = cur_bullets[-1] + b
                else:
                    cur_bullets.append(b)
    if cur_title is not None:
        md.discussions.append((cur_title, cur_bullets))

def _parse_chair_block(text: str, md: MeetingData):
    cur_person, cur_instrs = None, []
    for line in text.splitlines():
        if line.strip().startswith('|'):   # 跳過表格列
            continue
        if re.match(r'^\s*#{2,3}\s+', line):
            if cur_person is not None:
                md.chair_instrs.append((cur_person, cur_instrs[:]))
            cur_person = _clean(re.sub(r'^\s*#{2,3}\s+', '', line))
            cur_instrs = []
        elif cur_person is not None and re.match(r'^\s*[-*•]\s+', line):
            b = _clean(re.sub(r'^\s*[-*•]\s+', '', line))
            if b:
                cur_instrs.append(b)
        elif cur_person is not None and line.strip() and not re.match(r'^\s*#', line):
            b = _clean(line)
            if b and len(b) > 3:
                # 延續行偵測：合併未以句末符號結束的前一行（排除數字標記如 1. 2. 3.）
                _sent_end = re.search(
                    r'[。！？…!\?]\s*(?:\[\d+\])?\s*$'
                    r'|(?<!\d)\.\s*(?:\[\d+\])?\s*$',
                    cur_instrs[-1]) if cur_instrs else None
                if (cur_instrs
                        and not _sent_end
                        and not re.match(r'^\s*[-*•]', line)):
                    cur_instrs[-1] = cur_instrs[-1] + b
                else:
                    cur_instrs.append(b)
    if cur_person is not None:
        md.chair_instrs.append((cur_person, cur_instrs))

# ── 自由格式解析（fallback 補充）─────────────────────────────────────────────
def _parse_freeform(text: str, md: MeetingData):
    # 標題（# 一級標題 或第一行非空白、非表格、非 metadata 列）
    if not md.subject:
        for line in text.splitlines():
            stripped = line.strip()
            if not stripped or stripped.startswith('|'):
                continue   # 跳過空行和 Markdown 表格列
            if stripped.startswith('# '):
                md.subject = _clean(stripped[2:])
            elif not stripped.startswith('#'):
                md.subject = _clean(stripped)
            if md.subject:
                break

    # 日期
    if not md.date:
        dm = re.search(r'(\d{4}\s*年\s*\d+\s*月\s*\d+\s*日)', text)
        if dm:
            md.date = re.sub(r'\s+', '', dm.group(1))

    raw_secs = re.split(r'\n(?=##\s)', text)

    # Action Items
    if not md.action_items:
        for sec in raw_secs:
            first = sec.split('\n')[0].lower()
            if any(k in first for k in ('待辦', 'action', '行動', '待處理')):
                num = 1
                for line in sec.splitlines()[1:]:
                    # 表格行
                    parts = [p.strip() for p in line.split('|')]
                    parts = [p for p in parts if p and not re.match(r'^[-:＿]+$', p)]
                    if len(parts) >= 4 and re.match(r'^\d+', parts[0]):
                        md.action_items.append((parts[0], _clean(parts[1]),
                                                _clean(parts[2]), _clean(parts[3])))
                        continue
                    # 編號清單
                    m = re.match(r'\s*\d+\.\s+(.*)', line)
                    if m and m.group(1).strip():
                        md.action_items.append(
                            (str(num), _clean(m.group(1)), '待確認', '待確認'))
                        num += 1
                break

    # Discussion
    if not md.discussions:
        for sec in raw_secs:
            first = sec.split('\n')[0].lower()
            if any(k in first for k in ('主題', '討論', 'theme', 'discussion', '議題')):
                _parse_discussion_block(sec, md)
                if md.discussions:
                    break

    # Chair
    if not md.chair_instrs:
        for sec in raw_secs:
            first = sec.split('\n')[0]
            if re.search(r'(?:長|主管|總|長官|主席)\s*指示|指示\s*$|指示事項', first):
                _parse_chair_block(sec, md)
                break

# ── docx 建構 ─────────────────────────────────────────────────────────────────
def find_template() -> str | None:
    """找 meeting_sample/ 最新 .docx 作為範本。"""
    if not os.path.isdir(SAMPLE_DIR):
        return None
    files = sorted(
        [os.path.join(SAMPLE_DIR, f) for f in os.listdir(SAMPLE_DIR)
         if f.lower().endswith('.docx') and not f.startswith('~$')],
        key=os.path.getmtime, reverse=True)
    return files[0] if files else None

def build_docx(md: MeetingData, output_path: str) -> str:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    template = find_template()
    if not template:
        raise FileNotFoundError(
            f"找不到範本！請在 {SAMPLE_DIR} 放置至少一份集團格式 .docx")

    shutil.copy2(template, output_path)
    doc = Document(output_path)

    # ── Table 0：頁首資訊 ──────────────────────────────────────────────────────
    t0 = doc.tables[0]

    def _tc(row_idx, cell_idx):
        return t0.rows[row_idx]._tr.findall(qn('w:tc'))[cell_idx]

    _set_tc(qn, OxmlElement, _tc(0, 1), md.subject)
    _set_tc(qn, OxmlElement, _tc(1, 1), md.date)
    _set_tc(qn, OxmlElement, _tc(1, 3), md.time)
    _set_tc(qn, OxmlElement, _tc(1, 5), md.location)
    _set_tc(qn, OxmlElement, _tc(2, 1), md.attendees)
    _set_tc(qn, OxmlElement, _tc(3, 1), md.keeper)

    # ── Body para[1]：議程說明（支援以「；」分隔的多項議程）─────────────────────
    if len(doc.paragraphs) > 1:
        agenda_items = [a.strip() for a in md.agenda.replace('；', ';').split(';')
                        if a.strip()]
        if len(agenda_items) <= 1:
            # 單一行：直接寫入原段落
            _set_para(qn, OxmlElement, doc.paragraphs[1], md.agenda)
        else:
            # 多項：第一項寫入原段落，其餘複製段落插入其後
            import copy as _copy
            first_p = doc.paragraphs[1]._p
            _set_para(qn, OxmlElement, doc.paragraphs[1], agenda_items[0])
            ref_p = first_p   # 在此段落之後依序插入
            for item in agenda_items[1:]:
                new_p = _copy.deepcopy(first_p)
                # 清除 run 文字後寫入新內容
                for r in new_p.findall(qn('w:r')):
                    for t in r.findall(qn('w:t')):
                        t.text = ''
                if new_p.findall(qn('w:r')):
                    new_p.findall(qn('w:r'))[0].findall(qn('w:t'))[0].text = item
                ref_p.addnext(new_p)
                ref_p = new_p

    # ── Table 1：Action Items ──────────────────────────────────────────────────
    t1 = doc.tables[1]
    row_tmpl = copy.deepcopy(t1.rows[1]._tr)   # clone row 1 as template

    for row in list(t1.rows)[1:]:               # 刪除原始 data rows
        t1._tbl.remove(row._tr)

    for (num, text, dri, due) in md.action_items:
        new_tr = copy.deepcopy(row_tmpl)
        cells = new_tr.findall(qn('w:tc'))
        _set_tc(qn, OxmlElement, cells[0], num)
        _set_tc(qn, OxmlElement, cells[1], text)
        _set_tc(qn, OxmlElement, cells[2], dri)
        _set_tc(qn, OxmlElement, cells[3], due)
        # 強制 Action items 欄（cells[1]）靠左對齊，覆蓋範本的置中設定
        for p in cells[1].findall(qn('w:p')):
            pPr = p.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                p.insert(0, pPr)
            jc = pPr.find(qn('w:jc'))
            if jc is None:
                jc = OxmlElement('w:jc')
                pPr.append(jc)
            jc.set(qn('w:val'), 'left')
        t1._tbl.append(new_tr)

    # ── Body 段落：Discussion + Chair ─────────────────────────────────────────
    # 動態掃描 table1 之後的段落，找 bold（header）與 non-bold（bullet）模板
    # 不依賴硬式索引，避免不同範本結構造成錯誤
    body = doc.element.body
    body_kids = list(body)
    t1_pos = body_kids.index(t1._tbl)

    header_tmpl          = None
    bullet_tmpl          = None   # 優先：具有 ListParagraph 樣式的非粗體段落
    bullet_tmpl_fallback = None   # 備援：任意非粗體段落（不限樣式）
    for elem in body_kids[t1_pos + 1:]:
        if elem.tag != qn('w:p'):
            continue
        runs = elem.findall(qn('w:r'))
        if not runs:
            continue
        rpr = runs[0].find(qn('w:rPr'))
        is_bold = rpr is not None and rpr.find(qn('w:b')) is not None
        pPr    = elem.find(qn('w:pPr'))
        pStyle = pPr.find(qn('w:pStyle')) if pPr is not None else None
        style_val = (pStyle.get(qn('w:val'), '') if pStyle is not None else '').lower()
        is_list = 'list' in style_val   # ListParagraph / List Bullet / …
        if is_bold and header_tmpl is None:
            header_tmpl = copy.deepcopy(elem)
        elif not is_bold:
            if is_list and bullet_tmpl is None:
                bullet_tmpl = copy.deepcopy(elem)   # 優先選 List 樣式
            elif bullet_tmpl_fallback is None:
                bullet_tmpl_fallback = copy.deepcopy(elem)
        if header_tmpl is not None and bullet_tmpl is not None:
            break
    # 若無 List 樣式的非粗體段落，退而求其次
    if bullet_tmpl is None:
        bullet_tmpl = bullet_tmpl_fallback

    # Fallback：找任何段落
    if header_tmpl is None:
        for p in doc.paragraphs:
            if p._p.findall(qn('w:r')):
                header_tmpl = copy.deepcopy(p._p)
                break
    if bullet_tmpl is None:
        bullet_tmpl = copy.deepcopy(header_tmpl) if header_tmpl else OxmlElement('w:p')

    # 刪除 Table 1 之後的所有 <w:p>（保留 sectPr）
    for child in body_kids[t1_pos + 1:]:
        if child.tag != qn('w:sectPr'):
            body.remove(child)

    def _add(template_p, txt, bold: bool, is_bullet: bool = False,
             is_sub_header: bool = False):
        new_p = copy.deepcopy(template_p)
        # 清除所有 run 文字
        for r in list(new_p.findall(qn('w:r'))):
            new_p.remove(r)

        # ── 剝除 Markdown 粗體/斜體標記（** / * / ***），避免原始符號外露 ─────────
        import re as _re
        txt = _re.sub(r'\*{1,3}([^*]+)\*{1,3}', r'\1', txt)

        # ── sub_header：議題子標題，去自動編號、縮排 360、粗體 ──────────────────
        if is_sub_header:
            pPr = new_p.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                new_p.insert(0, pPr)
            numPr = pPr.find(qn('w:numPr'))
            if numPr is not None:
                pPr.remove(numPr)
            ind = pPr.find(qn('w:ind'))
            if ind is None:
                ind = OxmlElement('w:ind')
                pPr.append(ind)
            ind.set(qn('w:left'), '360')   # 議題子標題縮排（與 section header 區別）
            hanging_key = qn('w:hanging')
            if hanging_key in ind.attrib:
                del ind.attrib[hanging_key]

        # ── bullet 段落：移除 numPr 自動編號，改用縮排 ──────────────────────────
        elif is_bullet:
            pPr = new_p.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                new_p.insert(0, pPr)
            # 移除 numPr（否則繼承 header_tmpl 的中文編號樣式）
            numPr = pPr.find(qn('w:numPr'))
            if numPr is not None:
                pPr.remove(numPr)
            # left=720 twip：換行對齊位置
            # hanging=360 twip：「・ 」全形中點＋空白的實際寬度，讓第二行對齊第一行文字
            ind = pPr.find(qn('w:ind'))
            if ind is None:
                ind = OxmlElement('w:ind')
                pPr.append(ind)
            ind.set(qn('w:left'),    '720')
            ind.set(qn('w:hanging'), '360')
            # 加中點作為視覺 bullet 標記
            txt = '・ ' + txt

        # 取第一個 run 的 rPr 作為字型來源
        old_r = template_p.findall(qn('w:r'))
        rpr = None
        if old_r:
            rp = old_r[0].find(qn('w:rPr'))
            if rp is not None:
                rpr = copy.deepcopy(rp)
        # 建立新 run
        new_r = OxmlElement('w:r')
        if rpr is not None:
            # 修正 bold 旗標
            b_elem = rpr.find(qn('w:b'))
            bcs_elem = rpr.find(qn('w:bCs'))
            if bold:
                if b_elem is None:
                    rpr.insert(0, OxmlElement('w:b'))
                if bcs_elem is None:
                    rpr.append(OxmlElement('w:bCs'))
            else:
                if b_elem is not None:
                    rpr.remove(b_elem)
                if bcs_elem is not None:
                    rpr.remove(bcs_elem)
            new_r.append(rpr)
        t_elem = OxmlElement('w:t')
        t_elem.text = txt
        t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        new_r.append(t_elem)
        new_p.append(new_r)
        # 插入到 sectPr 之前
        sectPr = body.find(qn('w:sectPr'))
        if sectPr is not None:
            body.insert(list(body).index(sectPr), new_p)
        else:
            body.append(new_p)

    # 討論事項與決議
    _add(header_tmpl, '討論事項與決議', bold=True)
    for title, bullets in md.discussions:
        # 議題標題：子標題，去除自動編號、縮排 180 twip、保留粗體
        _add(header_tmpl, title, bold=True, is_sub_header=True)
        for b in bullets:
            _add(bullet_tmpl, b, bold=False, is_bullet=True)

    # 主管指示
    if md.chair_instrs:
        _add(header_tmpl, '主管指示', bold=True)   # 加一行 section header
        for person, instrs in md.chair_instrs:
            _add(header_tmpl, person + ('指示' if '指示' not in person else ''), bold=True, is_sub_header=True)
            for instr in instrs:
                _add(bullet_tmpl, instr, bold=False, is_bullet=True)  # 縮排但非編號

    doc.save(output_path)
    return output_path

# ── 輔助：設定 tc / para 文字（保留 rPr 字型）─────────────────────────────────
def _set_tc(qn, OxmlElement, tc_elem, text: str):
    """清除 <w:tc> 的 run，插入新文字，保留原始字型設定。"""
    paras = tc_elem.findall(qn('w:p'))
    if not paras:
        return
    _set_para_elem(qn, OxmlElement, paras[0], text)

def _set_para_elem(qn, OxmlElement, para_elem, text: str):
    runs = para_elem.findall(qn('w:r'))
    rpr = None
    for r in runs:
        rp = r.find(qn('w:rPr'))
        if rp is not None:
            rpr = copy.deepcopy(rp)
            break
    for r in runs:
        para_elem.remove(r)
    new_r = OxmlElement('w:r')
    if rpr is not None:
        new_r.append(rpr)
    t_elem = OxmlElement('w:t')
    t_elem.text = text
    t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    new_r.append(t_elem)
    para_elem.append(new_r)

def _set_para(qn, OxmlElement, para_obj, text: str):
    _set_para_elem(qn, OxmlElement, para_obj._p, text)

# ── HTML 輸出 ─────────────────────────────────────────────────────────────────
_ESC = str.maketrans({'&': '&amp;', '<': '&lt;', '>': '&gt;', '"': '&quot;'})

def _he(s: str) -> str:
    """HTML escape"""
    return s.translate(_ESC)

def _build_html_single(md: MeetingData) -> str:
    """仿 Word 版面，A4 比例，適合列印 / 封存。"""
    h = _he

    # ── Action Items rows ────────────────────────────────────────────────────
    if md.action_items:
        ai_rows = ''.join(
            f'<tr><td class="tc">{h(n)}</td><td>{h(t)}</td>'
            f'<td class="tc">{h(d)}</td><td class="tc">{h(u)}</td></tr>\n'
            for n, t, d, u in md.action_items)
    else:
        ai_rows = '<tr><td colspan="4" class="tc muted">（無待辦事項）</td></tr>\n'

    # ── Discussion ───────────────────────────────────────────────────────────
    disc_html = ''
    for title, bullets in md.discussions:
        disc_html += f'<p class="disc-title">{h(title)}</p>\n'
        disc_html += ''.join(f'<p class="bullet">・{h(b)}</p>\n' for b in bullets)

    # ── Chair Instructions ───────────────────────────────────────────────────
    chair_html = ''
    if md.chair_instrs:
        chair_html = '<p class="sec">主管指示</p>\n'
        for person, instrs in md.chair_instrs:
            chair_html += f'<p class="chair-person">{h(person)}</p>\n'
            chair_html += ''.join(
                f'<p class="bullet bold">・{h(i)}</p>\n' for i in instrs)

    return f'''<!DOCTYPE html>
<html lang="zh-TW">
<head>
<meta charset="UTF-8">
<title>{h(md.subject)} 會議紀錄</title>
<style>
body{{font-family:"標楷體","DFKai-SB","MingLiU",serif;background:#e8e8e8;
     margin:0;padding:32px 16px;font-size:12pt;color:#111;line-height:1.6}}
.page{{background:#fff;max-width:794px;margin:0 auto;padding:44px 52px;
      box-shadow:0 2px 14px rgba(0,0,0,.18)}}
h1{{text-align:center;font-size:16pt;letter-spacing:.15em;margin:0 0 24px}}
table{{width:100%;border-collapse:collapse;margin-bottom:18px}}
td{{border:1px solid #555;padding:5px 10px;vertical-align:middle}}
.lbl{{background:#d4d4d4;font-weight:bold;text-align:center;white-space:nowrap;
      width:1%}}
.hdr{{background:#4a4a4a;color:#fff;font-weight:bold;text-align:center}}
.tc{{text-align:center}}
.muted{{color:#999}}
.sec{{font-weight:bold;font-size:13pt;margin:20px 0 6px}}
.disc-title{{font-weight:bold;margin:10px 0 3px 6px}}
.bullet{{margin:3px 0 3px 28px}}
.bold{{font-weight:bold}}
.chair-person{{font-weight:bold;margin:10px 0 3px;border-left:3px solid #555;
              padding-left:8px}}
.agenda{{margin:0 0 20px 8px;color:#333}}
@media print{{
  body{{background:#fff;padding:0}}
  .page{{box-shadow:none;padding:24px 32px}}
  @page{{margin:1.5cm;size:A4}}
}}
</style>
</head>
<body>
<div class="page">
<h1>會&ensp;議&ensp;紀&ensp;錄</h1>

<table>
<tr><td class="lbl">會議名稱</td><td colspan="5">{h(md.subject)}</td></tr>
<tr>
  <td class="lbl">日&emsp;期</td><td>{h(md.date)}</td>
  <td class="lbl">時&emsp;間</td><td>{h(md.time)}</td>
  <td class="lbl">地&emsp;點</td><td>{h(md.location)}</td>
</tr>
<tr><td class="lbl">出&ensp;席&ensp;者</td><td colspan="5">{h(md.attendees)}</td></tr>
<tr><td class="lbl">記&ensp;錄&ensp;人</td><td colspan="5">{h(md.keeper)}</td></tr>
</table>

<p class="sec">會議議程</p>
<p class="agenda">一、{h(md.agenda)}</p>

<p class="sec">Action Items</p>
<table>
<tr>
  <td class="hdr" style="width:8%">項次</td>
  <td class="hdr">Action items</td>
  <td class="hdr" style="width:14%">DRI</td>
  <td class="hdr" style="width:18%">預計完成日</td>
</tr>
{ai_rows}
</table>

<p class="sec">討論事項與決議</p>
{disc_html}
{chair_html}
</div>
</body>
</html>'''


def _build_html_rwd(md: MeetingData) -> str:
    """響應式現代版面，行動裝置友善，純 CSS 無外部依賴。"""
    h = _he

    # ── Action Items rows ────────────────────────────────────────────────────
    if md.action_items:
        ai_rows = ''.join(
            f'<tr>'
            f'<td data-label="項次" class="tc">{h(n)}</td>'
            f'<td data-label="Action items">{h(t)}</td>'
            f'<td data-label="DRI" class="tc nowrap">{h(d)}</td>'
            f'<td data-label="預計完成日" class="tc nowrap">{h(u)}</td>'
            f'</tr>\n'
            for n, t, d, u in md.action_items)
    else:
        ai_rows = '<tr><td colspan="4" class="tc muted" style="padding:16px">（無待辦事項）</td></tr>\n'

    # ── Discussion ───────────────────────────────────────────────────────────
    disc_html = ''
    for title, bullets in md.discussions:
        bullets_html = ''.join(
            f'<div class="disc-bullet">{h(b)}</div>\n' for b in bullets)
        disc_html += f'''<div class="disc-item">
<div class="disc-title">{h(title)}</div>
{bullets_html}</div>\n'''

    # ── Chair Instructions ───────────────────────────────────────────────────
    chair_section = ''
    if md.chair_instrs:
        persons = ''
        for person, instrs in md.chair_instrs:
            instr_html = ''.join(
                f'<div class="chair-instr">▸ {h(i)}</div>\n' for i in instrs)
            persons += f'<div class="chair-person">🔴 {h(person)}</div>\n{instr_html}'
        chair_section = f'''<div class="card">
  <div class="card-hdr">主管指示</div>
  <div class="card-body">{persons}</div>
</div>\n'''

    # ── Info grid items ──────────────────────────────────────────────────────
    def info_item(label, value, wide=False):
        cls = 'info-item wide' if wide else 'info-item'
        return f'<div class="{cls}"><label>{h(label)}</label><span>{h(value)}</span></div>\n'

    info_grid = (
        info_item('日期', md.date) +
        info_item('時間', md.time) +
        info_item('地點', md.location) +
        info_item('出席者', md.attendees, wide=True) +
        info_item('記錄人', md.keeper, wide=True)
    )

    return f'''<!DOCTYPE html>
<html lang="zh-TW">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1.0">
<title>{h(md.subject)}</title>
<style>
:root{{--pri:#1a3a5c;--acc:#2e7d32;--red:#b71c1c;--bg:#f4f6f9;--card:#fff;
      --txt:#2c3e50;--dim:#607d8b;--bdr:#e0e0e0}}
*{{box-sizing:border-box;margin:0;padding:0}}
body{{font-family:system-ui,-apple-system,"Segoe UI","標楷體",sans-serif;
     background:var(--bg);color:var(--txt);line-height:1.65}}
/* ── Hero ─────────────────────────────────────────────────────────────── */
.hero{{background:var(--pri);color:#fff;padding:28px 20px 24px}}
.hero-badge{{display:inline-block;background:rgba(255,255,255,.18);
            border:1px solid rgba(255,255,255,.4);border-radius:20px;
            padding:2px 12px;font-size:.78rem;letter-spacing:.08em;margin-bottom:10px}}
.hero h1{{font-size:clamp(1.15rem,4vw,1.7rem);font-weight:700;margin-bottom:6px}}
.hero .meta{{font-size:.875rem;opacity:.8}}
/* ── Container & Card ─────────────────────────────────────────────────── */
.container{{max-width:900px;margin:0 auto;padding:20px 16px}}
.card{{background:var(--card);border-radius:10px;
      box-shadow:0 1px 8px rgba(0,0,0,.09);margin-bottom:18px;overflow:hidden}}
.card-hdr{{padding:10px 18px;font-weight:700;font-size:.92rem;letter-spacing:.04em;
           color:#fff}}
.card-hdr.green{{background:var(--acc)}}
.card-hdr.navy{{background:var(--pri)}}
.card-hdr.slate{{background:#37474f}}
.card-hdr.red{{background:var(--red)}}
.card-body{{padding:18px}}
/* ── Info Grid ────────────────────────────────────────────────────────── */
.info-grid{{display:grid;grid-template-columns:repeat(auto-fill,minmax(200px,1fr));gap:14px}}
.info-item label{{display:block;font-size:.72rem;font-weight:700;
                 text-transform:uppercase;color:var(--dim);margin-bottom:3px;
                 letter-spacing:.05em}}
.info-item span{{font-size:.95rem}}
.info-item.wide{{grid-column:1/-1}}
/* ── Action Table ─────────────────────────────────────────────────────── */
.action-wrap{{overflow-x:auto}}
table.ai{{width:100%;border-collapse:collapse;font-size:.9rem}}
.ai th{{background:#37474f;color:#fff;padding:9px 14px;text-align:left;
        white-space:nowrap}}
.ai td{{padding:9px 14px;border-bottom:1px solid var(--bdr)}}
.ai tr:last-child td{{border-bottom:none}}
.ai tr:nth-child(even) td{{background:#f8fafb}}
.tc{{text-align:center}}
.nowrap{{white-space:nowrap}}
.muted{{color:#aaa}}
/* ── Discussion ───────────────────────────────────────────────────────── */
.disc-item{{margin-bottom:18px}}
.disc-item:last-child{{margin-bottom:0}}
.disc-title{{font-weight:700;color:var(--pri);padding-left:12px;
            border-left:3px solid var(--acc);margin-bottom:8px;font-size:.98rem}}
.disc-bullet{{padding:3px 0 3px 22px;font-size:.93rem;position:relative}}
.disc-bullet::before{{content:"·";position:absolute;left:9px;color:var(--acc);
                     font-weight:bold}}
/* ── Chair ────────────────────────────────────────────────────────────── */
.chair-person{{font-weight:700;color:var(--red);margin:14px 0 6px;font-size:.98rem}}
.chair-person:first-child{{margin-top:0}}
.chair-instr{{padding:3px 0 3px 20px;font-size:.93rem;font-weight:600}}
/* ── Footer ───────────────────────────────────────────────────────────── */
.footer{{text-align:center;font-size:.78rem;color:var(--dim);padding:24px 0 32px}}
/* ── Mobile: stack table as cards ────────────────────────────────────── */
@media(max-width:540px){{
  .ai thead{{display:none}}
  .ai tr{{display:block;border:1px solid var(--bdr);border-radius:8px;
          margin-bottom:10px;overflow:hidden}}
  .ai td{{display:block;padding:7px 14px;border-bottom:1px solid #f0f0f0}}
  .ai td::before{{content:attr(data-label);display:block;
                 font-size:.72rem;font-weight:700;color:var(--dim);
                 text-transform:uppercase;margin-bottom:2px}}
  .ai tr:nth-child(even) td{{background:var(--card)}}
}}
@media print{{
  body{{background:#fff}}
  .card{{box-shadow:none;border:1px solid #ccc;page-break-inside:avoid}}
  .hero{{-webkit-print-color-adjust:exact;print-color-adjust:exact}}
}}
</style>
</head>
<body>

<div class="hero">
  <div class="container">
    <div class="hero-badge">集團會議紀錄</div>
    <h1>{h(md.subject)}</h1>
    <div class="meta">{h(md.date)}&emsp;{h(md.time)}&emsp;{h(md.location)}</div>
  </div>
</div>

<div class="container">

  <div class="card">
    <div class="card-hdr navy">會議基本資訊</div>
    <div class="card-body">
      <div class="info-grid">
{info_grid}      </div>
    </div>
  </div>

  <div class="card">
    <div class="card-hdr green">議程</div>
    <div class="card-body">{h(md.agenda)}</div>
  </div>

  <div class="card">
    <div class="card-hdr slate">Action Items</div>
    <div class="card-body action-wrap">
      <table class="ai">
        <thead>
          <tr>
            <th style="width:60px">項次</th>
            <th>Action items</th>
            <th style="width:110px">DRI</th>
            <th style="width:130px">預計完成日</th>
          </tr>
        </thead>
        <tbody>
{ai_rows}        </tbody>
      </table>
    </div>
  </div>

  <div class="card">
    <div class="card-hdr navy">討論事項與決議</div>
    <div class="card-body">
{disc_html}    </div>
  </div>

{chair_section}
  <div class="footer">
    由 NotebookLM 整理 &middot; {h(md.keeper)} &middot; {h(md.date)}
  </div>

</div>
</body>
</html>'''


def build_html(md: MeetingData, output_path: str, mode: str = 'single') -> str:
    """
    生成 HTML 格式會議紀錄並寫入檔案。
    mode='single' : 仿 Word A4 版面，適合列印/封存
    mode='rwd'    : 響應式現代版面，行動裝置友善
    回傳 output_path。
    """
    html = _build_html_rwd(md) if mode == 'rwd' else _build_html_single(md)
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(html)
    return output_path


# ── 公開入口 ──────────────────────────────────────────────────────────────────
def parse_summary(md: MeetingData, sections: dict | None = None) -> str:
    """回傳解析結果摘要，供 GUI log 診斷。"""
    sec_keys = ', '.join(sections.keys()) if sections else '—'
    return (
        f"subject='{md.subject}' | date='{md.date}' | time='{md.time}'\n"
        f"  attendees='{md.attendees[:30]}' | agenda='{md.agenda[:30]}'\n"
        f"  action_items={len(md.action_items)} | "
        f"discussions={len(md.discussions)} | chair_instrs={len(md.chair_instrs)}\n"
        f"  sections={sec_keys}"
    )


def convert(md_text: str, title_hint: str = '') -> list[str]:
    """
    將 MD 文字一次轉換為三種格式的集團格式會議紀錄。
    回傳 [docx_path, html_single_path, html_rwd_path, parse_summary_str]。
    各格式獨立 try，HTML 失敗不影響 Word 產出。
    """
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    sections = _split_sections(md_text)   # 保留給 parse_summary 診斷用
    md = parse_md(md_text)

    safe = re.sub(r'[\\/:*?"<>|]', '_',
                  md.subject or title_hint or '會議紀錄')[:40]
    ts   = datetime.datetime.now().strftime('%Y%m%d_%H%M')
    base = os.path.join(OUTPUT_DIR, f"{safe}_{ts}")

    docx_path   = base + '.docx'
    single_path = base + '_single.html'
    rwd_path    = base + '_rwd.html'
    errors      = []

    # 各格式獨立產出，互不影響
    try:
        build_docx(md, docx_path)
    except Exception as e:
        errors.append(f'Word: {e}')
        docx_path = ''

    try:
        build_html(md, single_path, mode='single')
    except Exception as e:
        errors.append(f'HTML-single: {e}')
        single_path = ''

    try:
        build_html(md, rwd_path, mode='rwd')
    except Exception as e:
        errors.append(f'HTML-rwd: {e}')
        rwd_path = ''

    summary = parse_summary(md, sections)
    if errors:
        summary += '\n  ⚠ errors: ' + ' | '.join(errors)

    return [docx_path, single_path, rwd_path, summary]

# ── CLI ───────────────────────────────────────────────────────────────────────
if __name__ == '__main__':
    ap = argparse.ArgumentParser(
        description='NLM MD → 集團格式會議紀錄（Word + HTML Single + HTML RWD）')
    ap.add_argument('input', nargs='?', help='.md 檔案路徑')
    ap.add_argument('--text', help='直接傳入 Markdown 字串')
    ap.add_argument('--clip', action='store_true', help='從剪貼簿讀取')
    ap.add_argument('--out', help='指定輸出主檔名前綴（不含副檔名）')
    args = ap.parse_args()

    if args.clip:
        try:
            import tkinter as tk
            root = tk.Tk(); root.withdraw()
            md_text = root.clipboard_get(); root.destroy()
        except Exception as e:
            print(f'無法讀取剪貼簿：{e}'); sys.exit(1)
    elif args.text:
        md_text = args.text
    elif args.input:
        md_text = Path(args.input).read_text(encoding='utf-8')
    else:
        print('請指定 .md 檔案、--text 或 --clip')
        ap.print_help(); sys.exit(1)

    if args.out:
        prefix = args.out
        os.makedirs(os.path.dirname(prefix) or '.', exist_ok=True)
        parsed = parse_md(md_text)
        outs = [
            build_docx(parsed, prefix + '.docx'),
            build_html(parsed, prefix + '_single.html', mode='single'),
            build_html(parsed, prefix + '_rwd.html',    mode='rwd'),
        ]
    else:
        outs = convert(md_text)

    for o in outs:
        print(f'✅  已生成：{o}')
    try:
        os.startfile(outs[0])   # 預設開啟 Word
    except Exception:
        pass
